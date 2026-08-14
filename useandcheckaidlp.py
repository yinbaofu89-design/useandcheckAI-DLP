import streamlit as st
import base64
import json
import uuid
import requests
import os
import mimetypes
from datetime import datetime
import time
import threading
import io
from PIL import Image
import re
import zipfile
from io import BytesIO

# ページ設定
st.set_page_config(
    page_title="DDS + AI 統合ツール",
    page_icon="🔒",
    layout="wide"
)

st.title("🔒 DDS + AI 統合コンテンツ検査ツール")
st.caption("DDSでポリシー違反をチェック後、AI APIで応答を生成")

# ==================== 初期化 ====================
if "messages" not in st.session_state:
    st.session_state.messages = []
if "info_check_results" not in st.session_state:
    st.session_state.info_check_results = []
if "txid" not in st.session_state:
    st.session_state.txid = str(uuid.uuid4())
if "filters" not in st.session_state:
    st.session_state.filters = [
        {"id": "c23de41e-f4a7-4b9e-9c1b-5b4eef283ec0", "name": "PCI"},
        {"id": "e58edfb6-bfa2-4256-ae28-ce929ba46bc8", "name": "source code detection"},
        {"id": "1443472b-c71f-49a0-bb44-06119fe48d0c", "name": "情報漏洩防止"}
    ]
if "dds_configured" not in st.session_state:
    st.session_state.dds_configured = False
if "ai_configured" not in st.session_state:
    st.session_state.ai_configured = False
if "uploaded_file_info" not in st.session_state:
    st.session_state.uploaded_file_info = None
if "file_checked" not in st.session_state:
    st.session_state.file_checked = False
if "file_violations" not in st.session_state:
    st.session_state.file_violations = []
if "file_approved" not in st.session_state:
    st.session_state.file_approved = False
if "file_data" not in st.session_state:
    st.session_state.file_data = None
if "filename" not in st.session_state:
    st.session_state.filename = None
if "ai_api_key" not in st.session_state:
    st.session_state.ai_api_key = "1234"
if "ai_api_url" not in st.session_state:
    st.session_state.ai_api_url = "http://localhost:1234/v1"
if "ai_model" not in st.session_state:
    st.session_state.ai_model = "Qwen3 8B - Q4_K_M"
if "selected_provider" not in st.session_state:
    st.session_state.selected_provider = "ローカル (LM Studio)"
if "debug_mode" not in st.session_state:
    st.session_state.debug_mode = False
if "debug_logs" not in st.session_state:
    st.session_state.debug_logs = []
if "log_counter" not in st.session_state:
    st.session_state.log_counter = 0
if "process_start_time" not in st.session_state:
    st.session_state.process_start_time = None
if "show_debug_panel" not in st.session_state:
    st.session_state.show_debug_panel = False
if "operation_mode" not in st.session_state:
    st.session_state.operation_mode = "Monitor"
if "send_target" not in st.session_state:
    st.session_state.send_target = "both"
if "info_check_done" not in st.session_state:
    st.session_state.info_check_done = False
if "show_dds_response" not in st.session_state:
    st.session_state.show_dds_response = False
if "blocked_content" not in st.session_state:
    st.session_state.blocked_content = False

# 情報チェックリスト項目
if "info_check_items" not in st.session_state:
    st.session_state.info_check_items = [
        {"name": "日本の電話番号情報", "enabled": True},
        {"name": "日本の住所情報", "enabled": True},
        {"name": "日本の名前情報", "enabled": True},
        {"name": "日本の銀行口座情報", "enabled": True},
        {"name": "日本のクレジットカード情報", "enabled": True},
        {"name": "日本のマイナンバー情報", "enabled": True},
        {"name": "メールアドレス", "enabled": True},
    ]

# ==================== AIプロバイダー設定 ====================
AI_PROVIDERS = {
    "DeepSeek": {
        "url": "https://api.deepseek.com/v1/chat/completions",
        "models": ["deepseek-chat", "deepseek-coder"],
        "default_model": "deepseek-chat",
        "api_key_required": True,
        "description": "DeepSeek API",
        "supports_file_upload": False
    },
    "OpenAI": {
        "url": "https://api.openai.com/v1/responses",
        "models": ["gpt-5", "gpt-5-mini"],
        "default_model": "gpt-5-mini",
        "api_key_required": True,
        "description": "OpenAI API",
        "supports_file_upload": True
    },
    "ローカル (LM Studio)": {
        "url": "http://localhost:1234/v1",
        "models": ["Qwen3 8B - Q4_K_M", "llama3-8b", "mistral-7b", "phi-3", "gemma-2b"],
        "default_model": "Qwen3 8B - Q4_K_M",
        "api_key_required": False,
        "default_api_key": "1234",
        "description": "LM Studio ローカルサーバー",
        "supports_file_upload": False
    },
    "ローカル (Ollama)": {
        "url": "http://localhost:11434/v1/chat/completions",
        "models": ["llama3", "mistral", "phi3", "gemma", "qwen", "llama2", "codellama", "llava"],
        "default_model": "llama3",
        "api_key_required": False,
        "default_api_key": "ollama",
        "description": "Ollama ローカルサーバー",
        "supports_file_upload": False
    },
    "Groq": {
        "url": "https://api.groq.com/openai/v1/chat/completions",
        "models": ["llama3-70b-8192", "llama3-8b-8192", "mixtral-8x7b-32768", "gemma2-9b-it"],
        "default_model": "llama3-70b-8192",
        "api_key_required": True,
        "description": "Groq API",
        "supports_file_upload": False
    }
}

# ==================== MIMEタイプ関数 ====================
def get_mime_type(filename):
    ext = os.path.splitext(filename)[1].lower()
    mime_map = {
        '.txt': 'text/plain', '.csv': 'text/csv', '.log': 'text/plain',
        '.ini': 'text/plain', '.cfg': 'text/plain', '.conf': 'text/plain',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.dot': 'application/msword', '.dotx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.template',
        '.docm': 'application/vnd.ms-word.document.macroEnabled.12',
        '.xls': 'application/vnd.ms-excel',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xlsm': 'application/vnd.ms-excel.sheet.macroEnabled.12',
        '.xlsb': 'application/vnd.ms-excel.sheet.binary.macroEnabled.12',
        '.ppt': 'application/vnd.ms-powerpoint',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        '.pptm': 'application/vnd.ms-powerpoint.presentation.macroEnabled.12',
        '.pps': 'application/vnd.ms-powerpoint', '.ppsx': 'application/vnd.openxmlformats-officedocument.presentationml.slideshow',
        '.pdf': 'application/pdf',
        '.eml': 'message/rfc822', '.msg': 'application/vnd.ms-outlook',
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png', '.gif': 'image/gif',
        '.bmp': 'image/bmp', '.tiff': 'image/tiff', '.tif': 'image/tiff',
        '.webp': 'image/webp', '.svg': 'image/svg+xml', '.ico': 'image/x-icon',
        '.zip': 'application/zip', '.7z': 'application/x-7z-compressed',
        '.rar': 'application/vnd.rar', '.tar': 'application/x-tar', '.gz': 'application/gzip',
        '.html': 'text/html', '.htm': 'text/html',
        '.xml': 'text/xml', '.json': 'application/json',
        '.css': 'text/css', '.js': 'application/javascript',
        '.rtf': 'application/rtf',
        '.odt': 'application/vnd.oasis.opendocument.text',
        '.ods': 'application/vnd.oasis.opendocument.spreadsheet',
        '.odp': 'application/vnd.oasis.opendocument.presentation',
    }
    return mime_map.get(ext, 'application/octet-stream')

def is_image_mime(mime_type):
    return bool(mime_type) and mime_type.lower().startswith("image/")

def build_data_url(file_bytes, mime_type):
    encoded = base64.b64encode(file_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{encoded}"

# ==================== ZIP作成関数（新規追加） ====================
def create_zip_for_dds(info_check_result, file_data, filename, user_message, send_target):
    """
    情報チェックAI分析結果と元のファイル/メッセージをZIPにまとめる
    
    Args:
        info_check_result: 情報チェックAIの分析結果（テキスト）
        file_data: 元のファイルデータ（バイナリ）
        filename: 元のファイル名
        user_message: 元のユーザーメッセージ
        send_target: 送信ターゲット（"file", "message", "both"）
    
    Returns:
        BytesIO: ZIPファイルのバイナリデータ
    """
    zip_buffer = BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # 1. 情報チェックAI分析結果を追加
        info_content = f"""=== 情報チェックAI分析結果 ===
{info_check_result}

=== 送信ターゲット ===
{send_target}

=== 作成日時 ===
{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        zip_file.writestr("info_check_result.txt", info_content.encode('utf-8'))
        
        # 2. 元のメッセージを追加（メッセージがある場合）
        if user_message and send_target in ["message", "both"]:
            message_content = f"""=== ユーザーメッセージ ===
{user_message}
"""
            zip_file.writestr("user_message.txt", message_content.encode('utf-8'))
        
        # 3. 元のファイルを追加（ファイルがある場合）
        if file_data and filename and send_target in ["file", "both"]:
            # 元のファイル名をそのまま使用
            zip_file.writestr(f"original_{filename}", file_data)
    
    zip_buffer.seek(0)
    return zip_buffer

# ==================== ZIPファイルをDDSに送信する関数（新規追加） ====================
def send_zip_to_dds(zip_data, dds_url, verify_ssl, source_type="zip", content_block_id=None):
    """
    ZIPファイルをDDSに送信する
    
    Args:
        zip_data: ZIPファイルのバイナリデータ（BytesIO）
        dds_url: DDSのURL
        verify_ssl: SSL検証フラグ
        source_type: 送信元タイプ
        content_block_id: コンテンツブロックID
    
    Returns:
        violations, request_id, response_data, error_info, elapsed
    """
    start_time = time.time()
    
    try:
        zip_data.seek(0)
        file_bytes = zip_data.read()
        b64_data = base64.b64encode(file_bytes).decode('utf-8')
        
        context = [
            {"name": "common.dataType", "value": ["DIM"]},
            {"name": "common.application", "value": [st.session_state.get("dlp_application", "securlet.box")]},
            {"name": "common.transactionId", "value": [st.session_state.txid]},
            {"name": "common.filter", "value": [f["id"] for f in st.session_state.filters]},
            {"name": "common.expectActionsAck", "value": ["true"]},
        ]
        
        if st.session_state.get("client_domain"):
            context.append({"name": "client.domain", "value": [st.session_state.client_domain]})
        if st.session_state.get("client_user"):
            context.append({"name": "client.user.id", "value": [st.session_state.client_user]})
        
        block_id = content_block_id or "zip-001"
        
        request_data = {
            "context": context,
            "subject": {
                "contentBlockId": "subject-001",
                "mimeType": "text/plain",
                "data": base64.b64encode(f"ZIPファイル: 情報チェックAI分析結果 + 元の内容".encode('utf-8')).decode('utf-8')
            },
            "attachments": [
                {
                    "contentBlockId": block_id,
                    "mimeType": "application/zip",
                    "data": b64_data,
                    "name": "info_check_package.zip"
                }
            ]
        }
        
        json_data = json.dumps(request_data, ensure_ascii=False)
        
        headers = {
            "Content-Type": "application/json",
            "Accept": "application/json"
        }
        
        if st.session_state.debug_mode:
            add_debug_log("DDS ZIPリクエスト", f"ZIPファイル送信 (サイズ: {len(file_bytes)}バイト)", "info", 0, {
                "url": dds_url,
                "zip_size": len(file_bytes),
                "request_data": request_data
            })
        
        response = requests.post(
            dds_url,
            data=json_data,
            headers=headers,
            verify=not verify_ssl,
            timeout=60
        )
        
        elapsed = time.time() - start_time
        
        try:
            response_json = response.json()
            
            if st.session_state.debug_mode:
                add_debug_log("DDS ZIPレスポンス", f"ステータス: {response.status_code}", 
                             "success" if response.status_code == 201 else "error", elapsed, {
                                "status_code": response.status_code,
                                "response": response_json
                            })
            
            if response.status_code == 201:
                violations = response_json.get("violation", [])
                if violations is None:
                    violations = []
                request_id = response_json.get("requestId")
                return violations, request_id, response_json, None, elapsed
            else:
                error_info = {
                    "status_code": response.status_code,
                    "response_text": response.text,
                    "headers": dict(response.headers)
                }
                return [], None, None, error_info, elapsed
                
        except Exception as e:
            error_info = {
                "status_code": response.status_code,
                "response_text": response.text,
                "error": str(e)
            }
            return [], None, None, error_info, elapsed
            
    except requests.exceptions.ConnectionError as e:
        error_info = {
            "error_type": "ConnectionError",
            "message": str(e),
            "dds_url": dds_url
        }
        return [], None, None, error_info, time.time() - start_time
    except requests.exceptions.Timeout as e:
        error_info = {
            "error_type": "Timeout",
            "message": str(e)
        }
        return [], None, None, error_info, time.time() - start_time
    except Exception as e:
        error_info = {
            "error_type": "Exception",
            "message": str(e)
        }
        import traceback
        error_info["traceback"] = traceback.format_exc()
        return [], None, None, error_info, time.time() - start_time

# ==================== デバッグログ関数 ====================
def add_debug_log(step, message, log_type="info", elapsed_time=None, details=None):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
    st.session_state.log_counter += 1
    log_entry = {
        "id": st.session_state.log_counter,
        "timestamp": timestamp,
        "step": step,
        "message": message,
        "type": log_type,
        "elapsed_time": elapsed_time,
        "details": details
    }
    st.session_state.debug_logs.append(log_entry)
    st.session_state.show_debug_panel = True

def clear_debug_logs():
    st.session_state.debug_logs = []
    st.session_state.log_counter = 0
    st.session_state.show_debug_panel = False

def render_debug_logs():
    if not st.session_state.debug_logs:
        st.info("📋 デバッグログはありません")
        return
    
    for log in st.session_state.debug_logs:
        log_id = log["id"]
        timestamp = log["timestamp"]
        step = log["step"]
        message = log["message"]
        log_type = log["type"]
        elapsed = log.get("elapsed_time")
        details = log.get("details")
        
        icon_map = {"info": "ℹ️", "success": "✅", "warning": "⚠️", "error": "❌"}
        color_map = {"info": "#0066cc", "success": "#00aa00", "warning": "#cc8800", "error": "#cc0000"}
        
        icon = icon_map.get(log_type, "ℹ️")
        color = color_map.get(log_type, "#0066cc")
        
        time_info = f"🕐 {timestamp}"
        if elapsed is not None:
            time_info += f" | ⏱️ {elapsed:.3f}秒"
        
        st.markdown(
            f"""
            <div style="border-left: 3px solid {color}; padding: 4px 8px; margin: 2px 0; 
                        background-color: #f8f9fa; border-radius: 3px; font-family: monospace; font-size: 12px;">
                <div><span style="font-weight: bold; color: {color};">[{step}]</span> {icon} {message}</div>
                <div style="font-size: 10px; color: #888;">{time_info}</div>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        if details:
            with st.expander(f"📝 詳細 (ID: {log_id})", expanded=False):
                st.json(details)
    
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ ログクリア", key="clear_logs_btn_panel", use_container_width=True):
            clear_debug_logs()
            st.rerun()
    with col2:
        if st.button("📋 コピー", key="copy_logs_btn_panel", use_container_width=True):
            log_text = "\n".join([
                f"[{log['timestamp']}] [{log['step']}] {log['message']} (経過: {log.get('elapsed_time', 0):.3f}秒)"
                for log in st.session_state.debug_logs
            ])
            st.code(log_text, language="text")
    
    st.caption(f"📊 ログ件数: {len(st.session_state.debug_logs)} 件")

# ==================== 情報チェックAIプロンプト ====================
def build_information_check_prompt(check_items):
    items = [x["name"] for x in check_items if x.get("enabled")]
    item_lines = "\n".join(f"- {item}" for item in items)

    return f"""あなたは情報漏えい防止(DLP)の補助分析を行うAIです。

【タスク】
以下のチェック項目に該当する情報が、提出されたファイルやメッセージに**含まれているかどうか**を判定してください。

【チェック項目】
{item_lines}

【出力ルール】
1. 該当する情報が**含まれている場合**、以下の形式で出力してください。
2. 該当する情報が**含まれていない場合**は、何も出力しないでください。
3. 推測や類推は行わないでください。
4. 説明や挨拶は一切不要です。

【出力形式】
[項目名]情報内包
ヒットした[項目名]内容-[実際の内容]

【出力例】
入力に「田中一郎」と「横浜市北区三丁目3-201」が含まれている場合：
名前情報内包
住所情報内包
ヒットした名前内容-田中一郎
ヒットした住所内容-横浜市北区三丁目3-201

【注意】
- 上記は出力形式の例です。実際に含まれている項目だけを出力してください。
- 含まれていない項目は絶対に出力しないでください。"""

# ==================== ファイル内容抽出関数 ====================
def extract_file_content(file_bytes, filename, max_chars=5000):
    """ファイルからテキスト内容を抽出する"""
    mime = get_mime_type(filename)
    
    # テキストファイル
    if mime.startswith("text/") or filename.lower().endswith((".txt", ".csv", ".log", ".json", ".xml", ".html", ".htm", ".md")):
        try:
            decoded = file_bytes.decode("utf-8", errors="ignore")
            if len(decoded) > max_chars:
                return decoded[:max_chars] + f"\n... (省略: {len(decoded) - max_chars}文字)"
            return decoded
        except:
            return f"（バイナリファイル: {len(file_bytes)}バイト）"
    
    # Word文書 (.docx)
    elif filename.lower().endswith(".docx"):
        try:
            try:
                import docx
            except ImportError:
                return f"（Wordファイル: {filename}、{len(file_bytes)}バイト - python-docxが必要です）"
            
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([para.text for para in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            if len(text) > max_chars:
                return text[:max_chars] + f"\n... (省略: {len(text) - max_chars}文字)"
            return text
        except Exception as e:
            return f"（Word抽出エラー: {e}）"
    
    # 古いWord文書 (.doc)
    elif filename.lower().endswith(".doc"):
        return f"（Wordファイル: {filename}、{len(file_bytes)}バイト - .doc形式はpython-docx非対応）"
    
    # PDFファイル
    elif filename.lower().endswith(".pdf"):
        try:
            try:
                import PyPDF2
            except ImportError:
                return f"（PDFファイル: {filename}、{len(file_bytes)}バイト - PyPDF2が必要です）"
            
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            if len(text) > max_chars:
                return text[:max_chars] + f"\n... (省略: {len(text) - max_chars}文字)"
            return text
        except Exception as e:
            return f"（PDF抽出エラー: {e}）"
    
    # Excelファイル (.xlsx)
    elif filename.lower().endswith((".xlsx", ".xls")):
        try:
            try:
                import openpyxl
            except ImportError:
                return f"（Excelファイル: {filename}、{len(file_bytes)}バイト - openpyxlが必要です）"
            
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            text = ""
            for sheet in wb.worksheets:
                text += f"【シート: {sheet.title}】\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) for cell in row if cell is not None])
                    if row_text:
                        text += row_text + "\n"
            if len(text) > max_chars:
                return text[:max_chars] + f"\n... (省略: {len(text) - max_chars}文字)"
            return text
        except Exception as e:
            return f"（Excel抽出エラー: {e}）"
    
    # その他のバイナリファイル
    else:
        return f"（バイナリファイル: {filename}、{len(file_bytes)}バイト）"

# ==================== 情報チェックAI実行 ====================
def run_information_check_ai(file_bytes, filename, user_message, check_items, additional_context=None):
    prompt = build_information_check_prompt(check_items)
    
    messages = [
        {"role": "system", "content": prompt},
        {"role": "user", "content": "以下の内容を分析し、該当する情報があれば出力してください。"}
    ]
    
    user_content = ""
    if user_message:
        user_content += f"【メッセージ】\n{user_message}\n\n"
    
    if file_bytes and filename:
        user_content += f"【ファイル: {filename}】\n"
        file_text = extract_file_content(file_bytes, filename)
        user_content += file_text
        user_content += "\n"
    
    if additional_context:
        user_content += "\n【追加コンテキスト情報】\n"
        user_content += additional_context
    
    if user_content:
        messages.append({"role": "user", "content": user_content})
    else:
        messages.append({"role": "user", "content": "分析する内容がありません。"})
    
    return call_ai_api(messages, max_tokens=500)

# ==================== AI API呼び出し ====================
def normalize_api_url(url):
    url = url.strip()
    if url.endswith('/v1'):
        return url + '/chat/completions'
    if url.endswith('/v1/'):
        return url + 'chat/completions'
    if not url.endswith('/chat/completions') and not url.endswith('/completions'):
        if not url.endswith('/'):
            return url + '/v1/chat/completions'
        else:
            return url + 'v1/chat/completions'
    return url

def call_ai_api(messages, max_tokens=200):
    start_time = time.time()
    
    try:
        api_key = st.session_state.ai_api_key
        api_url = normalize_api_url(st.session_state.ai_api_url)
        model_name = st.session_state.ai_model
        
        if not api_url or not model_name:
            st.error("AI設定が不完全です。")
            return None, time.time() - start_time
        
        provider = st.session_state.selected_provider
        if provider in AI_PROVIDERS and AI_PROVIDERS[provider].get("api_key_required", True) and not api_key:
            st.error(f"{provider}はAPI Keyが必須です。")
            return None, time.time() - start_time
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": model_name,
            "messages": messages,
            "max_tokens": max_tokens,
            "temperature": 0.3,
            "stream": False
        }
        
        if st.session_state.debug_mode:
            add_debug_log("AIリクエスト", f"AI API呼び出し (モデル: {model_name})", "info", 0, {
                "url": api_url,
                "model": model_name,
                "temperature": 0.3,
                "messages": [{"role": m["role"], "content": str(m["content"])[:200]} for m in messages]
            })
        
        response = requests.post(api_url, headers=headers, json=data, timeout=120)
        
        elapsed = time.time() - start_time
        
        if response.status_code == 200:
            result = response.json()
            
            if st.session_state.debug_mode:
                add_debug_log("AIレスポンス", f"AI API応答受信", "success", elapsed, {
                    "response": result
                })
            
            content = None
            
            if "choices" in result and len(result["choices"]) > 0:
                choice = result["choices"][0]
                message = choice.get("message", {})
                content = message.get("content", "")
                
                if not content or content.strip() == "":
                    content = "検出結果なし"
                
                if content and content != "検出結果なし":
                    lines = content.strip().split('\n')
                    filtered_lines = []
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        skip_patterns = ['はい、', 'わかりました', '以下は', '分析結果', '【', '（', '注意', '重要', '確認します']
                        should_skip = False
                        for pattern in skip_patterns:
                            if line.startswith(pattern):
                                should_skip = True
                                break
                        if not should_skip:
                            filtered_lines.append(line)
                    
                    if filtered_lines:
                        content = '\n'.join(filtered_lines)
                    else:
                        content = "検出結果なし"
                
                return content, elapsed
            
            elif "response" in result:
                return result["response"], elapsed
            
            else:
                st.error(f"予期しないレスポンス形式: {result}")
                return None, elapsed
                
        else:
            error_msg = f"AI APIエラー: {response.status_code} - {response.text}"
            st.error(error_msg)
            add_debug_log("AIエラー", error_msg, "error", elapsed)
            return None, elapsed
            
    except requests.exceptions.ConnectionError as e:
        error_msg = f"❌ 接続エラー: APIサーバー ({api_url}) に接続できませんでした"
        st.error(error_msg)
        add_debug_log("AIエラー", error_msg, "error", time.time() - start_time)
        return None, time.time() - start_time
    except requests.exceptions.Timeout:
        error_msg = "❌ タイムアウト: サーバーからの応答がありませんでした"
        st.error(error_msg)
        add_debug_log("AIエラー", error_msg, "error", time.time() - start_time)
        return None, time.time() - start_time
    except Exception as e:
        error_msg = f"AI API呼び出しエラー: {e}"
        st.error(error_msg)
        add_debug_log("AIエラー", error_msg, "error", time.time() - start_time)
        return None, time.time() - start_time

# ==================== OpenAIファイル送信 ====================
def call_openai_with_file(messages, file_bytes=None, filename=None, mime_type=None, max_tokens=200):
    start_time = time.time()
    try:
        api_key = st.session_state.ai_api_key
        model_name = st.session_state.ai_model
        if not api_key or not model_name:
            st.error("OpenAI API Keyまたはモデルが設定されていません。")
            return None, time.time() - start_time

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        input_items = []
        for msg in messages:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            input_items.append({
                "role": role,
                "content": [{"type": "input_text", "text": str(content)}]
            })

        if file_bytes is not None and filename:
            mime = mime_type or get_mime_type(filename)
            if is_image_mime(mime):
                image_url = build_data_url(file_bytes, mime)
                input_items.append({
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": f"添付ファイル: {filename}"},
                        {"type": "input_image", "image_url": image_url}
                    ]
                })
            else:
                upload_headers = {"Authorization": f"Bearer {api_key}"}
                files = {
                    "file": (filename, file_bytes, mime)
                }
                upload_data = {"purpose": "user_data"}
                upload_response = requests.post(
                    "https://api.openai.com/v1/files",
                    headers=upload_headers,
                    data=upload_data,
                    files=files,
                    timeout=120
                )
                if upload_response.status_code not in (200, 201):
                    st.error(f"OpenAI Files APIエラー: {upload_response.status_code} - {upload_response.text}")
                    return None, time.time() - start_time

                file_id = upload_response.json().get("id")
                if not file_id:
                    st.error("OpenAI Files APIからfile_idを取得できませんでした。")
                    return None, time.time() - start_time

                input_items.append({
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": f"添付ファイル {filename} を読み取り、ユーザーの質問に答えてください。"},
                        {"type": "input_file", "file_id": file_id}
                    ]
                })

        payload = {
            "model": model_name,
            "input": input_items,
            "max_output_tokens": max_tokens
        }

        response = requests.post(
            "https://api.openai.com/v1/responses",
            headers=headers,
            json=payload,
            timeout=120
        )
        elapsed = time.time() - start_time

        if response.status_code != 200:
            st.error(f"OpenAI Responses APIエラー: {response.status_code} - {response.text}")
            return None, elapsed

        result = response.json()
        output_text = result.get("output_text")
        if output_text:
            return output_text, elapsed

        texts = []
        for item in result.get("output", []):
            for content in item.get("content", []):
                if content.get("type") == "output_text" and content.get("text"):
                    texts.append(content["text"])
        return ("\n".join(texts) if texts else "検出結果なし"), elapsed

    except requests.exceptions.Timeout:
        st.error("❌ OpenAI APIタイムアウト")
        return None, time.time() - start_time
    except requests.exceptions.ConnectionError as e:
        st.error(f"❌ OpenAI API接続エラー: {e}")
        return None, time.time() - start_time
    except Exception as e:
        st.error(f"❌ OpenAIファイル送信エラー: {e}")
        return None, time.time() - start_time

# ==================== AIファイル送信（統合） ====================
def call_ai_with_file(messages, file_bytes=None, filename=None, mime_type=None, max_tokens=200):
    provider = st.session_state.selected_provider
    
    if provider == "OpenAI":
        return call_openai_with_file(messages, file_bytes, filename, mime_type, max_tokens)
    else:
        if file_bytes and filename:
            file_text = extract_file_content(file_bytes, filename)
            for i, msg in enumerate(messages):
                if msg.get("role") == "user":
                    messages[i] = {
                        "role": "user",
                        "content": msg.get("content", "") + f"\n\n【ファイル: {filename}】\n{file_text}"
                    }
                    break
        return call_ai_api(messages, max_tokens)

# ==================== DDS送信関数（従来のテキスト送信用） ====================
def send_detection_request(file_obj, source_type, dds_url, verify_ssl, data_type="DIM", content_block_id=None, mime_type=None):
    start_time = time.time()
    
    try:
        file_obj.seek(0)
        file_bytes = file_obj.read()
        b64_data = base64.b64encode(file_bytes).decode('utf-8')
        
        context = [
            {"name": "common.dataType", "value": [data_type]},
            {"name": "common.application", "value": [st.session_state.get("dlp_application", "securlet.box")]},
            {"name": "common.transactionId", "value": [st.session_state.txid]},
            {"name": "common.filter", "value": [f["id"] for f in st.session_state.filters]},
            {"name": "common.expectActionsAck", "value": ["true"]},
        ]
        
        if st.session_state.get("client_domain"):
            context.append({"name": "client.domain", "value": [st.session_state.client_domain]})
        if st.session_state.get("client_user"):
            context.append({"name": "client.user.id", "value": [st.session_state.client_user]})
        
        if source_type == "file":
            file_mime = get_mime_type(file_obj.name)
            if hasattr(file_obj, 'type') and file_obj.type and file_obj.type != 'application/octet-stream':
                file_mime = file_obj.type
            
            block_id = re.sub(r'[^a-zA-Z0-9-]', '-', file_obj.name) + "-001"
            
            request_data = {
                "context": context,
                "subject": {
                    "contentBlockId": "subject-001",
                    "mimeType": "text/plain",
                    "data": base64.b64encode(f"ファイル: {file_obj.name}".encode('utf-8')).decode('utf-8')
                },
                "attachments": [
                    {
                        "contentBlockId": block_id,
                        "mimeType": file_mime,
                        "data": b64_data,
                        "name": file_obj.name
                    }
                ]
            }
            
        else:
            block_id = content_block_id or "message-001"
            mime = mime_type or "text/plain"
            
            request_data = {
                "context": context,
                "subject": {
                    "contentBlockId": block_id,
                    "mimeType": "text/plain",
                    "data": b64_data
                }
            }
        
        json_data = json.dumps(request_data, ensure_ascii=False)
        
        headers = {
            "Content-Type": "application/json",
            "Accept": "application/json"
        }
        
        if st.session_state.debug_mode:
            add_debug_log("DDSリクエスト", f"リクエスト送信 (タイプ: {source_type})", "info", 0, {
                "url": dds_url,
                "source_type": source_type,
                "data_type": data_type,
                "request_data": request_data
            })
        
        response = requests.post(
            dds_url,
            data=json_data,
            headers=headers,
            verify=not verify_ssl,
            timeout=60
        )
        
        elapsed = time.time() - start_time
        
        try:
            response_json = response.json()
            
            if st.session_state.debug_mode:
                add_debug_log("DDSレスポンス", f"ステータス: {response.status_code}", 
                             "success" if response.status_code == 201 else "error", elapsed, {
                                "status_code": response.status_code,
                                "response": response_json
                            })
            
            if response.status_code == 201:
                violations = response_json.get("violation", [])
                if violations is None:
                    violations = []
                request_id = response_json.get("requestId")
                return violations, request_id, response_json, None, elapsed
            else:
                error_info = {
                    "status_code": response.status_code,
                    "response_text": response.text,
                    "headers": dict(response.headers)
                }
                return [], None, None, error_info, elapsed
                
        except Exception as e:
            error_info = {
                "status_code": response.status_code,
                "response_text": response.text,
                "error": str(e)
            }
            return [], None, None, error_info, elapsed
            
    except requests.exceptions.ConnectionError as e:
        error_info = {
            "error_type": "ConnectionError",
            "message": str(e),
            "dds_url": dds_url
        }
        return [], None, None, error_info, time.time() - start_time
    except requests.exceptions.Timeout as e:
        error_info = {
            "error_type": "Timeout",
            "message": str(e)
        }
        return [], None, None, error_info, time.time() - start_time
    except Exception as e:
        error_info = {
            "error_type": "Exception",
            "message": str(e)
        }
        import traceback
        error_info["traceback"] = traceback.format_exc()
        return [], None, None, error_info, time.time() - start_time

# ==================== MessageWrapperクラス ====================
class MessageWrapper:
    def __init__(self, content, name):
        self.content = content
        self.name = name
        self.type = "text/plain"
        self.size = len(content)
    def read(self):
        return self.content.encode('utf-8')
    def seek(self, pos):
        pass

class BytesWrapper:
    def __init__(self, data, name, mime_type="application/octet-stream"):
        self.data = data
        self.name = name
        self.type = mime_type
        self.size = len(data)
    def read(self):
        return self.data
    def seek(self, pos):
        pass

# ==================== コンテンツ取得関数 ====================
def get_content_for_dds(file_data, filename, user_message, send_target):
    content = ""
    if send_target in ["message", "both"] and user_message:
        content += f"【メッセージ】\n{user_message}\n"
    if send_target in ["file", "both"] and file_data and filename:
        content += f"【ファイル: {filename}】\n"
        file_text = extract_file_content(file_data, filename)
        content += file_text
    return content

# ==================== 履歴クリア関数 ====================
def clear_conversation():
    st.session_state.messages = []
    st.session_state.info_check_results = []
    st.session_state.txid = str(uuid.uuid4())
    st.session_state.blocked_content = False

# ==================== DDSレスポンス整形関数 ====================
def format_dds_response(response_data, violations, request_id):
    if not response_data:
        return "（DDSレスポンスなし）"
    
    lines = []
    lines.append("📋 **DDSレスポンス詳細**")
    lines.append("")
    
    if request_id:
        lines.append(f"🆔 Request ID: `{request_id}`")
    
    if violations and len(violations) > 0:
        lines.append(f"⚠️ 違反件数: {len(violations)}件")
        lines.append("")
        for i, v in enumerate(violations, 1):
            v_name = v.get("name", "不明")
            v_id = v.get("id", "不明")
            v_severity = v.get("severity", "不明")
            lines.append(f"  {i}. **{v_name}**")
            lines.append(f"     - ID: `{v_id}`")
            lines.append(f"     - 重大度: {v_severity}")
    else:
        lines.append("✅ 違反なし")
    
    lines.append("")
    lines.append("---")
    lines.append("📄 **JSONレスポンス:**")
    lines.append("```json")
    lines.append(json.dumps(response_data, ensure_ascii=False, indent=2))
    lines.append("```")
    
    return "\n".join(lines)

# ==================== AI回答をDDSでチェックする関数 ====================
def check_ai_response_with_dds(ai_response, step_name):
    start_time = time.time()
    
    add_debug_log(f"{step_name}-AI回答チェック", "AI回答をDDSに送信して検査中...", "info")
    
    ai_message_wrapper = MessageWrapper(ai_response, "ai_response")
    violations, request_id, response_data, error_info, elapsed = send_detection_request(
        ai_message_wrapper,
        "message",
        st.session_state.dds_url,
        st.session_state.verify_ssl,
        data_type="DIM",
        content_block_id="ai-response-check-001"
    )
    
    if violations is None:
        violations = []
    
    add_debug_log(f"{step_name}-AI回答チェック", f"AI回答DDS検査完了 (違反: {len(violations)}件)", 
                 "warning" if violations else "success", elapsed, response_data)
    
    return violations, request_id, response_data, error_info, elapsed

# ==================== サイドバー設定 ====================
with st.sidebar:
    st.header("⚙️ 設定")
    
    st.subheader("👤 クライアント情報")
    client_domain = st.text_input(
        "ドメイン",
        value=st.session_state.get("client_domain", "tds-cnx.com"),
        key="client_domain_input"
    )
    client_user = st.text_input(
        "ユーザー名",
        value=st.session_state.get("client_user", "Baofu"),
        key="client_user_input"
    )
    dlp_application = st.text_input(
        "Application",
        value=st.session_state.get("dlp_application", "securlet.box"),
        key="dlp_application_input"
    )
    st.session_state.client_domain = client_domain
    st.session_state.client_user = client_user
    st.session_state.dlp_application = dlp_application

    st.divider()

    st.subheader("🔄 動作モード")
    operation_mode = st.radio(
        "モード選択",
        ["Monitor", "Inline"],
        index=0 if st.session_state.operation_mode == "Monitor" else 1,
        horizontal=True
    )
    st.session_state.operation_mode = operation_mode
    
    st.subheader("📤 送信ターゲット")
    send_target = st.radio(
        "送信内容",
        ["ファイルのみ", "メッセージのみ", "両方"],
        index=2,
        horizontal=True
    )
    target_map = {"ファイルのみ": "file", "メッセージのみ": "message", "両方": "both"}
    st.session_state.send_target = target_map[send_target]

    st.divider()

    st.subheader("📋 DDSレスポンス表示")
    show_dds = st.checkbox(
        "DDSレスポンスをチャット履歴に表示",
        value=st.session_state.get("show_dds_response", False),
        help="チェックすると、すべてのDDSレスポンス詳細がチャット履歴に表示されます"
    )
    st.session_state.show_dds_response = show_dds
    
    if show_dds:
        st.info("✅ DDSレスポンス表示: ON")
    else:
        st.info("🔘 DDSレスポンス表示: OFF")

    st.divider()

    st.subheader("🔎 情報チェックAI")
    st.caption("AIによるDLP補助分析。初期状態では全項目ONです。")

    for idx, item in enumerate(st.session_state.info_check_items):
        item["enabled"] = st.checkbox(
            item["name"],
            value=item.get("enabled", True),
            key=f"info_check_{idx}"
        )

    with st.expander("➕ 検知項目を追加"):
        new_info_item = st.text_input(
            "追加する検知内容",
            key="new_info_item",
            placeholder="例：社内機密コード"
        )
        if st.button("追加", key="add_info_item_btn", use_container_width=True):
            if new_info_item.strip():
                st.session_state.info_check_items.append(
                    {"name": new_info_item.strip(), "enabled": True}
                )
                st.rerun()

    st.divider()

    st.subheader("🔍 DDS設定")
    dds_host = st.text_input(
        "DDSサーバーIP",
        value=st.session_state.get("dds_host", "20.89.66.42"),
        key="dds_host_input"
    )
    st.session_state.dds_host = dds_host
    dds_port = st.text_input("ポート", value="443")
    use_ssl = st.checkbox("SSL/TLSを使用", value=False)
    protocol = "https" if use_ssl else "http"
    dds_url = f"{protocol}://{dds_host}:{dds_port}/v2.0/DetectionRequests"
    st.session_state.dds_url = dds_url
    st.session_state.verify_ssl = use_ssl
    
    st.divider()
    
    st.subheader("🤖 AI設定")
    
    provider_options = list(AI_PROVIDERS.keys())
    default_index = provider_options.index("ローカル (LM Studio)") if "ローカル (LM Studio)" in provider_options else 0
    
    selected_provider = st.selectbox(
        "AIプロバイダー",
        provider_options,
        index=default_index
    )
    st.session_state.selected_provider = selected_provider
    
    provider_config = AI_PROVIDERS[selected_provider]
    st.caption(f"📌 {provider_config.get('description', '')}")
    
    st.session_state.ai_api_url = provider_config["url"]
    st.info(f"📡 API URL: {provider_config['url']}")
    
    model_options = provider_config["models"].copy() if provider_config["models"] else []
    model_options.append("その他 (カスタム)")
    
    current_model = st.session_state.ai_model
    if current_model not in model_options:
        current_display = "その他 (カスタム)"
    else:
        current_display = current_model
    
    selected_model = st.selectbox(
        "モデル",
        model_options,
        index=model_options.index(current_display) if current_display in model_options else 0
    )
    
    if selected_model == "その他 (カスタム)":
        custom_model = st.text_input(
            "カスタムモデル名",
            value=st.session_state.ai_model if st.session_state.ai_model not in provider_config["models"] else "",
            placeholder="モデル名を自由に入力してください"
        )
        if custom_model:
            st.session_state.ai_model = custom_model
    else:
        st.session_state.ai_model = selected_model
    
    if provider_config.get("api_key_required", True):
        st.session_state.ai_api_key = st.text_input(
            "API Key",
            value=st.session_state.ai_api_key,
            type="password",
            placeholder="APIキーを入力してください"
        )
    else:
        default_key = provider_config.get("default_api_key", "")
        st.session_state.ai_api_key = default_key
        st.info(f"🔑 API Key: {default_key} (自動設定)")
    
    st.divider()
    
    st.subheader("🐛 デバッグ設定")
    debug_mode = st.checkbox(
        "デバッグモードを有効にする",
        value=st.session_state.debug_mode
    )
    st.session_state.debug_mode = debug_mode
    
    if st.session_state.debug_logs:
        show_panel = st.checkbox(
            "📋 デバッグパネルを表示",
            value=st.session_state.get("show_debug_panel", True)
        )
        st.session_state.show_debug_panel = show_panel
    else:
        st.info("📋 デバッグログがありません")
    
    st.divider()
    
    if st.button("🔗 AI接続テスト", use_container_width=True):
        if st.session_state.ai_api_url and st.session_state.ai_model:
            test_result, elapsed = call_ai_api([
                {"role": "user", "content": "Hello, this is a test. Please respond with 'OK'."}
            ], max_tokens=10)
            if test_result:
                st.success(f"✅ AI接続成功！ (モデル: {st.session_state.ai_model})")
                st.caption(f"⏱️ 応答時間: {elapsed:.2f}秒")
                st.session_state.ai_configured = True
            else:
                st.error("❌ AI接続失敗。設定を確認してください。")
        else:
            st.warning("⚠️ API URLとモデル名を確認してください")
    
    st.divider()
    
    st.subheader("📋 検出フィルター")
    st.caption("最大10個")
    
    filters_to_remove = []
    for i, f in enumerate(st.session_state.filters):
        cols = st.columns([3, 1])
        with cols[0]:
            st.text_input(
                f"フィルター {i+1}",
                value=f["id"],
                key=f"filter_{i}",
                label_visibility="collapsed"
            )
            st.caption(f"📌 {f.get('name', '')}")
        with cols[1]:
            if st.button("🗑️", key=f"remove_filter_{i}"):
                filters_to_remove.append(i)
    
    for idx in sorted(filters_to_remove, reverse=True):
        st.session_state.filters.pop(idx)
        st.rerun()
    
    if len(st.session_state.filters) < 10:
        with st.expander("➕ フィルターを追加"):
            new_filter_id = st.text_input("フィルターID (GUID)", key="new_filter_id")
            new_filter_name = st.text_input("フィルター名", key="new_filter_name")
            if st.button("追加", key="add_filter_btn", use_container_width=True):
                if new_filter_id:
                    st.session_state.filters.append({
                        "id": new_filter_id,
                        "name": new_filter_name or f"フィルター {len(st.session_state.filters)+1}"
                    })
                    st.rerun()
    
    st.divider()
    
    st.text_input("トランザクションID", value=st.session_state.txid, disabled=True)
    if st.button("🔄 新しいIDを生成", key="generate_txid_btn", use_container_width=True):
        st.session_state.txid = str(uuid.uuid4())
        st.rerun()
    
    st.divider()
    
    if st.button("🗑️ 会話履歴をクリア", key="clear_conversation_btn", use_container_width=True):
        clear_conversation()
        st.rerun()
    
    st.divider()
    st.caption(f"📡 プロバイダー: **{st.session_state.get('selected_provider', '未設定')}**")
    st.caption(f"🤖 モデル: **{st.session_state.get('ai_model', '未設定')}**")

# ==================== メインコンテンツ ====================
show_debug = st.session_state.get("show_debug_panel", False) and len(st.session_state.debug_logs) > 0

if show_debug:
    main_col, debug_col = st.columns([2, 1])
else:
    main_col = st.container()
    debug_col = None

with main_col:
    if st.session_state.debug_logs:
        st.info(f"📋 デバッグログ: {len(st.session_state.debug_logs)}件")
    
    mode_color = "🟢" if st.session_state.operation_mode == "Monitor" else "🔵"
    st.caption(f"{mode_color} 現在のモード: **{st.session_state.operation_mode}** | 送信ターゲット: **{send_target}**")
    
    if st.session_state.operation_mode == "Inline":
        st.info("🔒 **Inlineモード**: 送信内容をDDSでチェックし、違反があればブロックします。その後、情報チェックAI分析を実行します。")
    
    if st.session_state.show_dds_response:
        st.info("📋 DDSレスポンス表示: **ON** - すべてのDDSレスポンスが表示されます")
    
    chat_container = st.container()
    with chat_container:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.write(message["content"])
                if "violations" in message and message["violations"]:
                    st.error(f"🚫 ポリシー違反: {', '.join([v['name'] for v in message['violations']])}")
                if "dds_response" in message and message["dds_response"] and st.session_state.show_dds_response:
                    with st.expander("📋 DDSレスポンス詳細", expanded=False):
                        st.markdown(message["dds_response"])
                if "ai_check_dds_response" in message and message["ai_check_dds_response"] and st.session_state.show_dds_response:
                    with st.expander("📋 AI回答DDSチェック詳細", expanded=False):
                        st.markdown(message["ai_check_dds_response"])

    st.subheader("📎 ファイルアップロード")
    
    uploaded_file = st.file_uploader(
        "ファイル選択",
        type=[".txt", ".csv", ".log", ".doc", ".docx", ".xls", ".xlsx", 
              ".ppt", ".pptx", ".pdf", ".eml", ".msg",
              ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".svg", ".ico",
              ".zip", ".7z", ".rar", ".tar", ".gz",
              ".html", ".htm", ".xml", ".json", ".css", ".js",
              ".rtf", ".odt", ".ods", ".odp"],
        key="file_uploader",
        label_visibility="collapsed"
    )

    if uploaded_file:
        if st.session_state.filename != uploaded_file.name:
            st.session_state.file_checked = False
            st.session_state.file_violations = []
            st.session_state.file_approved = False
            st.session_state.filename = uploaded_file.name
            uploaded_file.seek(0)
            st.session_state.file_data = uploaded_file.read()
            st.session_state.file_checked = False
        
        file_mime = get_mime_type(uploaded_file.name)
        if uploaded_file.type and uploaded_file.type != 'application/octet-stream':
            file_mime = uploaded_file.type
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.info(f"📎 {st.session_state.filename}")
        with col2:
            st.info(f"📊 {len(st.session_state.file_data)/1024:.1f} KB")
        with col3:
            st.info(f"📝 {file_mime}")
        with col4:
            ext = st.session_state.filename.split('.')[-1].upper() if '.' in st.session_state.filename else "不明"
            st.info(f"📄 {ext}")
        
        if file_mime and file_mime.startswith('image/'):
            try:
                image = Image.open(io.BytesIO(st.session_state.file_data))
                st.image(image, caption=f"画像プレビュー: {st.session_state.filename}", width=300)
            except:
                pass
        
        elif file_mime and "text" in file_mime:
            try:
                content = st.session_state.file_data.decode('utf-8', errors='ignore')
                st.text_area("📄 ファイル内容プレビュー", content[:1000], height=150)
                if len(content) > 1000:
                    st.caption(f"... 他 {len(content) - 1000} 文字")
            except:
                pass

    st.divider()

    user_input = st.chat_input("メッセージを入力してください...")

    # ==================== 送信処理 ====================
    if user_input:
        st.session_state.process_start_time = time.time()
        st.session_state.blocked_content = False
        
        with st.chat_message("user"):
            st.write(user_input)
            if uploaded_file:
                st.write(f"📎 {st.session_state.filename}")
        
        file_data = st.session_state.file_data if uploaded_file else None
        filename = st.session_state.filename if uploaded_file else None
        send_target = st.session_state.send_target
        
        if send_target == "message" and not user_input:
            st.error("❌ メッセージが入力されていません")
            st.stop()
        if send_target == "file" and not file_data:
            st.error("❌ ファイルがアップロードされていません")
            st.stop()
        if send_target == "both" and not user_input and not file_data:
            st.error("❌ メッセージとファイルの両方がありません")
            st.stop()
        
        # ==================== Monitorモード ====================
        if st.session_state.operation_mode == "Monitor":
            st.info(f"📊 **Monitorモード**で実行します")
            st.info("🔄 1つ目のリクエスト: 元の内容 → DDS → AI → AI回答DDSチェック → DDS再検査")
            st.info("🔄 2つ目のリクエスト: 情報チェックAI（ZIPでDDS送信）")
            
            # ============================================================
            # 1つ目のリクエスト（元の内容）
            # ============================================================
            
            # ---- ステップ1: 元の内容をDDSに送信 ----
            step1_start = time.time()
            add_debug_log("Monitor-1-ステップ1", "元の内容をDDSに送信して検査中...", "info")
            status1 = st.info("🔍 [Monitor-1-1] 元の内容をDDSに送信して検査中...")
            
            content_for_dds = get_content_for_dds(file_data, filename, user_input, send_target)
            if not content_for_dds.strip():
                st.error("❌ 送信するコンテンツがありません")
                st.stop()
            
            content_wrapper = MessageWrapper(content_for_dds, "content")
            violations, request_id, response_data, error_info, elapsed1 = send_detection_request(
                content_wrapper,
                "message",
                st.session_state.dds_url,
                st.session_state.verify_ssl,
                data_type="DIM",
                content_block_id="content-001"
            )
            
            if violations is None:
                violations = []
            
            status1.empty()
            st.info(f"✅ [Monitor-1-1] DDS検査完了 (⏱️ {elapsed1:.2f}秒)")
            add_debug_log("Monitor-1-ステップ1", f"DDS検査完了", "success", elapsed1, response_data)
            
            if error_info:
                st.error(f"❌ DDSエラー: {error_info}")
                add_debug_log("Monitor-1-ステップ1", f"エラー: {error_info}", "error", elapsed1)
                st.stop()
            
            if violations and len(violations) > 0:
                policy_names = [v["name"] for v in violations]
                st.warning(f"⚠️ {len(violations)}件のポリシー違反: {', '.join(policy_names)}")
                add_debug_log("Monitor-1-ステップ1", f"ポリシー違反: {', '.join(policy_names)}", "warning", elapsed1)
            else:
                st.success("✅ ポリシー違反はありませんでした")
                add_debug_log("Monitor-1-ステップ1", "ポリシー違反なし", "success", elapsed1)
            
            # ---- ステップ2: AIに送信 ----
            add_debug_log("Monitor-1-ステップ2", "AIにリクエストを送信中...", "info")
            st.info("🤖 [Monitor-1-2] AIにリクエストを送信中...")
            
            ai_messages = []
            for msg in st.session_state.messages:
                if msg["role"] != "assistant" or "violations" not in msg:
                    ai_messages.append({"role": msg["role"], "content": msg["content"]})
            
            current_msg = user_input
            ai_messages.append({"role": "user", "content": current_msg})
            
            start_time = time.time()
            status_placeholder = st.empty()
            status_placeholder.info(f"⏳ AI応答を待っています...")
            
            with st.spinner(f"🤖 {st.session_state.ai_model} に送信中..."):
                if file_data and filename:
                    ai_response, ai_elapsed = call_ai_with_file(
                        ai_messages,
                        file_bytes=file_data,
                        filename=filename,
                        mime_type=get_mime_type(filename),
                        max_tokens=200
                    )
                else:
                    ai_response, ai_elapsed = call_ai_api(ai_messages, max_tokens=200)
            
            status_placeholder.empty()
            elapsed_time = time.time() - start_time
            
            if not ai_response:
                st.error(f"❌ AIからの応答がありませんでした")
                add_debug_log("Monitor-1-ステップ2", f"AI応答なし", "error", ai_elapsed if 'ai_elapsed' in locals() else 0)
                st.stop()
            
            st.success(f"✅ [Monitor-1-2] AIからレスポンスを受信 (⏱️ {ai_elapsed:.2f}秒)")
            add_debug_log("Monitor-1-ステップ2", f"AI応答受信", "success", ai_elapsed)
            
            # ---- ステップ2.5: AI回答をDDSでチェック ----
            st.info("🔍 [Monitor-1-2.5] AI回答をDDSでチェック中...")
            
            ai_check_violations, ai_check_request_id, ai_check_response_data, ai_check_error, ai_check_elapsed = check_ai_response_with_dds(
                ai_response, "Monitor-1"
            )
            
            st.info(f"✅ [Monitor-1-2.5] AI回答DDSチェック完了 (⏱️ {ai_check_elapsed:.2f}秒)")
            
            ai_check_dds_response = format_dds_response(ai_check_response_data, ai_check_violations, ai_check_request_id) if st.session_state.show_dds_response else None
            ai_response_has_violation = ai_check_violations and len(ai_check_violations) > 0
            
            # ---- ステップ3: AI応答をDDSに再送信（再検査） ----
            step3_start = time.time()
            add_debug_log("Monitor-1-ステップ3", "AI応答をDDSに送信して再検査中...", "info")
            st.info("🔍 [Monitor-1-3] AI応答をDDSに送信して再検査中...")
            
            ai_message_wrapper = MessageWrapper(ai_response, "ai_response")
            v2, request_id2, response_data2, error_info2, elapsed3 = send_detection_request(
                ai_message_wrapper,
                "message",
                st.session_state.dds_url,
                st.session_state.verify_ssl,
                data_type="DIM",
                content_block_id="ai-response-001"
            )
            
            if v2 is None:
                v2 = []
            
            st.info(f"✅ [Monitor-1-3] DDS再検査完了 (⏱️ {elapsed3:.2f}秒)")
            add_debug_log("Monitor-1-ステップ3", f"DDS再検査完了", "success", elapsed3, response_data2)
            
            if error_info2:
                st.error(f"❌ DDSエラー: {error_info2}")
                add_debug_log("Monitor-1-ステップ3", f"エラー: {error_info2}", "error", elapsed3)
                st.stop()
            
            dds_response_text = format_dds_response(response_data2, v2, request_id2) if st.session_state.show_dds_response else None
            
            # ---- ステップ4: 結果表示 ----
            if ai_response_has_violation:
                ai_policy_names = [v["name"] for v in ai_check_violations]
                st.warning(f"⚠️ AIの回答にポリシー違反が検出されました: {', '.join(ai_policy_names)}")
                add_debug_log("Monitor-1-ステップ2.5", f"AI回答にポリシー違反: {', '.join(ai_policy_names)}", "warning", ai_check_elapsed)
            
            if v2 and len(v2) > 0:
                policy_names = [v["name"] for v in v2]
                st.warning(f"⚠️ AI応答に{len(v2)}件のポリシー違反: {', '.join(policy_names)}")
                add_debug_log("Monitor-1-ステップ4", f"AI応答にポリシー違反: {', '.join(policy_names)}", "warning", elapsed3)
                
                error_msg = f"⚠️ AIの回答にDLP違反が検出されました（違反: {', '.join(policy_names)}）\n\n---\n{ai_response}"
                if ai_response_has_violation:
                    error_msg = f"⚠️ AIの回答にDLP違反が検出されました（違反: {', '.join(ai_policy_names)}）\n\n---\n{ai_response}"
                
                msg_data = {
                    "role": "assistant",
                    "content": error_msg,
                    "violations": v2
                }
                if dds_response_text:
                    msg_data["dds_response"] = dds_response_text
                if ai_check_dds_response:
                    msg_data["ai_check_dds_response"] = ai_check_dds_response
                st.session_state.messages.append(msg_data)
                
                with st.chat_message("assistant"):
                    st.warning(f"⚠️ AIの回答にDLP違反が検出されました")
                    st.write(ai_response)
                    st.caption(f"⏱️ 応答時間: {elapsed_time:.2f}秒")
                    if dds_response_text:
                        with st.expander("📋 DDSレスポンス詳細", expanded=False):
                            st.markdown(dds_response_text)
                    if ai_check_dds_response:
                        with st.expander("📋 AI回答DDSチェック詳細", expanded=False):
                            st.markdown(ai_check_dds_response)
            else:
                st.success("✅ AI応答にポリシー違反はありません")
                add_debug_log("Monitor-1-ステップ4", "AI応答にポリシー違反なし", "success", elapsed3)
                
                msg_data = {
                    "role": "assistant",
                    "content": ai_response,
                    "violations": []
                }
                if dds_response_text:
                    msg_data["dds_response"] = dds_response_text
                if ai_check_dds_response:
                    msg_data["ai_check_dds_response"] = ai_check_dds_response
                st.session_state.messages.append(msg_data)
                
                with st.chat_message("assistant"):
                    st.write(ai_response)
                    st.caption(f"⏱️ 応答時間: {elapsed_time:.2f}秒")
                    if dds_response_text:
                        with st.expander("📋 DDSレスポンス詳細", expanded=False):
                            st.markdown(dds_response_text)
                    if ai_check_dds_response:
                        with st.expander("📋 AI回答DDSチェック詳細", expanded=False):
                            st.markdown(ai_check_dds_response)
            
            # ============================================================
            # 2つ目のリクエスト: 情報チェックAI（ZIPでDDS送信）
            # ============================================================
            st.divider()
            st.info("🔄 2つ目のリクエスト: 情報チェックAI（ZIPでDDS送信）")
            
            # ---- ステップ5: 情報チェックAIを実行 ----
            info_items_enabled = [x for x in st.session_state.info_check_items if x.get("enabled")]
            info_result = None
            info_elapsed = 0
            
            if info_items_enabled:
                add_debug_log("Monitor-2-ステップ1", "情報チェックAIで分析中...", "info")
                st.info("🔍 [Monitor-2-1] 情報チェックAIで分析中...")
                
                additional_context = ""
                if ai_response_has_violation:
                    ai_policy_names = [v["name"] for v in ai_check_violations]
                    additional_context += f"【AI回答のDDSチェック結果】\n"
                    additional_context += f"- 違反検出: あり\n"
                    additional_context += f"- 違反ポリシー: {', '.join(ai_policy_names)}\n"
                    additional_context += f"- 違反件数: {len(ai_check_violations)}件\n"
                    additional_context += f"- チェックID: {ai_check_request_id}\n"
                else:
                    additional_context += f"【AI回答のDDSチェック結果】\n"
                    additional_context += f"- 違反検出: なし\n"
                    additional_context += f"- チェックID: {ai_check_request_id}\n"
                
                if v2 and len(v2) > 0:
                    v2_policy_names = [v["name"] for v in v2]
                    additional_context += f"\n【AI回答のDDS再検査結果】\n"
                    additional_context += f"- 違反検出: あり\n"
                    additional_context += f"- 違反ポリシー: {', '.join(v2_policy_names)}\n"
                    additional_context += f"- 違反件数: {len(v2)}件\n"
                else:
                    additional_context += f"\n【AI回答のDDS再検査結果】\n"
                    additional_context += f"- 違反検出: なし\n"
                
                info_result, info_elapsed = run_information_check_ai(
                    file_data,
                    filename,
                    user_input,
                    info_items_enabled,
                    additional_context=additional_context
                )
                
                if info_result:
                    st.success(f"✅ [Monitor-2-1] 情報チェックAI分析完了 (⏱️ {info_elapsed:.2f}秒)")
                    add_debug_log(
                        "Monitor-2-ステップ1",
                        "情報チェックAI分析完了",
                        "success",
                        info_elapsed,
                        {
                            "check_items": [x["name"] for x in info_items_enabled],
                            "response": info_result
                        }
                    )
                    with st.expander("🔎 情報チェックAI分析結果", expanded=True):
                        st.code(info_result, language="text")
                else:
                    st.warning("⚠️ 情報チェックAIの分析結果を取得できませんでした")
                    add_debug_log("Monitor-2-ステップ1", "分析結果を取得できませんでした", "error", info_elapsed)
            
            # ---- ステップ6: ZIP作成してDDSに送信（新規） ----
            if info_result and info_result != "検出結果なし":
                step6_start = time.time()
                add_debug_log("Monitor-2-ステップ2", "ZIPファイルを作成してDDSに送信中...", "info")
                st.info("📦 [Monitor-2-2] ZIPファイルを作成してDDSに送信中...")
                
                # ZIP作成
                zip_data = create_zip_for_dds(
                    info_check_result=info_result,
                    file_data=file_data,
                    filename=filename,
                    user_message=user_input,
                    send_target=send_target
                )
                
                zip_size = len(zip_data.getvalue())
                st.info(f"📦 ZIP作成完了: {zip_size/1024:.1f} KB")
                
                # ZIPをDDSに送信
                v3, request_id3, response_data3, error_info3, elapsed6 = send_zip_to_dds(
                    zip_data=zip_data,
                    dds_url=st.session_state.dds_url,
                    verify_ssl=st.session_state.verify_ssl,
                    source_type="zip",
                    content_block_id="info-check-zip-001"
                )
                
                if v3 is None:
                    v3 = []
                
                st.info(f"✅ [Monitor-2-2] DDS ZIP検査完了 (⏱️ {elapsed6:.2f}秒)")
                add_debug_log("Monitor-2-ステップ2", f"ZIP DDS検査完了", "success", elapsed6, response_data3)
                
                if error_info3:
                    st.error(f"❌ DDSエラー: {error_info3}")
                    add_debug_log("Monitor-2-ステップ2", f"エラー: {error_info3}", "error", elapsed6)
                else:
                    info_dds_response = format_dds_response(response_data3, v3, request_id3) if st.session_state.show_dds_response else None
                    
                    if v3 and len(v3) > 0:
                        policy_names = [v["name"] for v in v3]
                        st.warning(f"⚠️ 情報チェックAI分析結果（ZIP）に{len(v3)}件のポリシー違反: {', '.join(policy_names)}")
                        add_debug_log("Monitor-2-ステップ3", f"ZIPにポリシー違反: {', '.join(policy_names)}", "warning", elapsed6)
                    else:
                        st.success("✅ 情報チェックAI分析結果（ZIP）にポリシー違反はありません")
                        add_debug_log("Monitor-2-ステップ3", "ZIPにポリシー違反なし", "success", elapsed6)
                    
                    # ZIP内の内容を説明するメッセージ
                    zip_contents = []
                    zip_contents.append("📦 **ZIPファイル内容:**")
                    zip_contents.append(f"  - info_check_result.txt: 情報チェックAI分析結果")
                    if user_input and send_target in ["message", "both"]:
                        zip_contents.append(f"  - user_message.txt: ユーザーメッセージ")
                    if file_data and filename and send_target in ["file", "both"]:
                        zip_contents.append(f"  - original_{filename}: 元のファイル")
                    zip_contents.append("")
                    
                    info_msg = f"📊 **情報チェックAI分析結果（ZIPでDDS送信）**\n\n"
                    info_msg += "\n".join(zip_contents)
                    info_msg += "\n"
                    if v3 and len(v3) > 0:
                        policy_names = [v["name"] for v in v3]
                        info_msg += f"⚠️ 検出された違反: {', '.join(policy_names)}\n\n"
                    else:
                        info_msg += "✅ 違反は検出されませんでした\n\n"
                    info_msg += "---\n"
                    info_msg += f"**分析結果:**\n```\n{info_result}\n```"
                    
                    msg_data = {
                        "role": "assistant",
                        "content": info_msg,
                        "violations": v3,
                        "is_info_check": True
                    }
                    if info_dds_response:
                        msg_data["dds_response"] = info_dds_response
                    st.session_state.messages.append(msg_data)
                    
                    with st.chat_message("assistant"):
                        st.markdown(info_msg)
                        if info_dds_response:
                            with st.expander("📋 DDSレスポンス詳細", expanded=False):
                                st.markdown(info_dds_response)
            elif info_result == "検出結果なし":
                st.info("ℹ️ 情報チェックAIで検出結果がなかったため、ZIP送信をスキップしました")
            else:
                st.warning("⚠️ 情報チェックAIの結果がないため、ZIP送信をスキップしました")
            
            st.success(f"✅ **Monitorモード完了** (合計: {time.time() - st.session_state.process_start_time:.2f}秒)")
            add_debug_log("Monitor-完了", f"全ての処理が完了しました", "success", time.time() - st.session_state.process_start_time)
        
        # ==================== Inlineモード ====================
        else:
            st.info(f"📊 **Inlineモード**で実行します")
            st.info("🔒 送信内容をDDSでチェックし、違反があればブロックします")
            st.info("🔄 情報チェックAI（ZIPでDDS送信）")
            
            # ============================================================
            # ステップ1: 元の内容をDDSに送信（ブロックチェック）
            # ============================================================
            step1_start = time.time()
            add_debug_log("Inline-ステップ1", "元の内容をDDSに送信してブロックチェック...", "info")
            status1 = st.info("🔍 [Inline-1] 元の内容をDDSに送信してブロックチェック中...")
            
            content_for_dds = get_content_for_dds(file_data, filename, user_input, send_target)
            if not content_for_dds.strip():
                st.error("❌ 送信するコンテンツがありません")
                st.stop()
            
            content_wrapper = MessageWrapper(content_for_dds, "content")
            violations, request_id, response_data, error_info, elapsed1 = send_detection_request(
                content_wrapper,
                "message",
                st.session_state.dds_url,
                st.session_state.verify_ssl,
                data_type="DIM",
                content_block_id="content-001"
            )
            
            if violations is None:
                violations = []
            
            status1.empty()
            st.info(f"✅ [Inline-1] DDSブロックチェック完了 (⏱️ {elapsed1:.2f}秒)")
            add_debug_log("Inline-ステップ1", f"DDSブロックチェック完了", "success", elapsed1, response_data)
            
            if error_info:
                st.error(f"❌ DDSエラー: {error_info}")
                add_debug_log("Inline-ステップ1", f"エラー: {error_info}", "error", elapsed1)
                st.stop()
            
            # ---- ブロック判定 ----
            content_has_violation = violations and len(violations) > 0
            ai_response = None
            ai_response_has_violation = False
            ai_check_violations = []
            ai_check_request_id = None
            ai_check_dds_response = None
            ai_elapsed = 0
            elapsed_time = 0
            
            if content_has_violation:
                policy_names = [v["name"] for v in violations]
                st.session_state.blocked_content = True
                
                st.error(f"""
                🚫 **【ブロック】ポリシー違反が検出されました**
                
                あなたの送信した内容に以下のポリシー違反が含まれているため、送信をブロックしました。
                
                **違反ポリシー:** {', '.join(policy_names)}
                
                この内容はAIに送信されませんでした。
                """)
                
                add_debug_log("Inline-ステップ1", f"ブロック: ポリシー違反 {', '.join(policy_names)}", "error", elapsed1)
                
                block_msg = f"""
🚫 **【ブロック】ポリシー違反が検出されました**

あなたの送信した内容に以下のポリシー違反が含まれているため、送信をブロックしました。

**違反ポリシー:** {', '.join(policy_names)}

この内容はAIに送信されませんでした。
"""
                msg_data = {
                    "role": "assistant",
                    "content": block_msg,
                    "violations": violations
                }
                if st.session_state.show_dds_response:
                    dds_response_text = format_dds_response(response_data, violations, request_id)
                    msg_data["dds_response"] = dds_response_text
                st.session_state.messages.append(msg_data)
                
                with st.chat_message("assistant"):
                    st.error(block_msg)
                    if st.session_state.show_dds_response and dds_response_text:
                        with st.expander("📋 DDSレスポンス詳細", expanded=False):
                            st.markdown(dds_response_text)
                
                st.info("🔄 ブロックされましたが、情報チェックAI分析を継続します...")
            else:
                st.success("✅ ポリシー違反はありません - 継続処理します")
                add_debug_log("Inline-ステップ1", "ポリシー違反なし - 継続", "success", elapsed1)
                
                # ---- ブロックされていない場合のみAIに送信 ----
                add_debug_log("Inline-ステップ1.5", "AIにリクエストを送信中...", "info")
                st.info("🤖 [Inline-1.5] AIにリクエストを送信中...")
                
                ai_messages = []
                for msg in st.session_state.messages:
                    if msg["role"] != "assistant" or "violations" not in msg:
                        ai_messages.append({"role": msg["role"], "content": msg["content"]})
                
                current_msg = user_input
                ai_messages.append({"role": "user", "content": current_msg})
                
                start_time = time.time()
                status_placeholder = st.empty()
                status_placeholder.info(f"⏳ AI応答を待っています...")
                
                with st.spinner(f"🤖 {st.session_state.ai_model} に送信中..."):
                    if file_data and filename:
                        ai_response, ai_elapsed = call_ai_with_file(
                            ai_messages,
                            file_bytes=file_data,
                            filename=filename,
                            mime_type=get_mime_type(filename),
                            max_tokens=200
                        )
                    else:
                        ai_response, ai_elapsed = call_ai_api(ai_messages, max_tokens=200)
                
                status_placeholder.empty()
                elapsed_time = time.time() - start_time
                
                if not ai_response:
                    st.error(f"❌ AIからの応答がありませんでした")
                    add_debug_log("Inline-ステップ1.5", f"AI応答なし", "error", ai_elapsed if 'ai_elapsed' in locals() else 0)
                    st.stop()
                
                st.success(f"✅ [Inline-1.5] AIからレスポンスを受信 (⏱️ {ai_elapsed:.2f}秒)")
                add_debug_log("Inline-ステップ1.5", f"AI応答受信", "success", ai_elapsed)
                
                # ---- ステップ1.6: AI回答をDDSでチェック ----
                st.info("🔍 [Inline-1.6] AI回答をDDSでチェック中...")
                
                ai_check_violations, ai_check_request_id, ai_check_response_data, ai_check_error, ai_check_elapsed = check_ai_response_with_dds(
                    ai_response, "Inline"
                )
                
                st.info(f"✅ [Inline-1.6] AI回答DDSチェック完了 (⏱️ {ai_check_elapsed:.2f}秒)")
                
                ai_check_dds_response = format_dds_response(ai_check_response_data, ai_check_violations, ai_check_request_id) if st.session_state.show_dds_response else None
                ai_response_has_violation = ai_check_violations and len(ai_check_violations) > 0
                
                if ai_response_has_violation:
                    ai_policy_names = [v["name"] for v in ai_check_violations]
                    st.error(f"""
                    🚨 **【警告】AIの回答にポリシー違反が検出されました**
                    
                    AIが生成した回答に以下のポリシー違反が含まれているため、回答を表示しません。
                    
                    **違反ポリシー:** {', '.join(ai_policy_names)}
                    
                    この回答はDLPの補助情報として記録されます。
                    """)
                    
                    add_debug_log("Inline-ステップ1.6", f"AI回答にポリシー違反: {', '.join(ai_policy_names)} - 表示ブロック", "error", ai_check_elapsed)
                    
                    warn_msg = f"""
🚨 **【警告】AIの回答にポリシー違反が検出されました**

AIが生成した回答に以下のポリシー違反が含まれているため、回答を表示しません。

**違反ポリシー:** {', '.join(ai_policy_names)}

この回答はDLPの補助情報として記録されます。
"""
                    msg_data = {
                        "role": "assistant",
                        "content": warn_msg,
                        "violations": ai_check_violations,
                        "ai_response_blocked": True
                    }
                    if ai_check_dds_response:
                        msg_data["dds_response"] = ai_check_dds_response
                    st.session_state.messages.append(msg_data)
                    
                    with st.chat_message("assistant"):
                        st.error(warn_msg)
                        if ai_check_dds_response:
                            with st.expander("📋 DDSレスポンス詳細", expanded=False):
                                st.markdown(ai_check_dds_response)
                else:
                    st.success("✅ AI回答にポリシー違反はありません")
                    add_debug_log("Inline-ステップ1.6", "AI回答にポリシー違反なし", "success", ai_check_elapsed)
                    
                    msg_data = {
                        "role": "assistant",
                        "content": ai_response,
                        "violations": []
                    }
                    if ai_check_dds_response:
                        msg_data["dds_response"] = ai_check_dds_response
                    st.session_state.messages.append(msg_data)
                    
                    with st.chat_message("assistant"):
                        st.write(ai_response)
                        st.caption(f"⏱️ 応答時間: {elapsed_time:.2f}秒")
                        if ai_check_dds_response:
                            with st.expander("📋 DDSレスポンス詳細", expanded=False):
                                st.markdown(ai_check_dds_response)
            
            # ============================================================
            # ステップ2: 情報チェックAI（ZIPでDDS送信）
            # ============================================================
            info_items_enabled = [x for x in st.session_state.info_check_items if x.get("enabled")]
            info_result = None
            info_elapsed = 0
            
            if info_items_enabled:
                add_debug_log("Inline-ステップ2", "情報チェックAIで分析中...", "info")
                st.info("🔍 [Inline-2] 情報チェックAIで分析中...")
                
                additional_context = ""
                if ai_response and not content_has_violation:
                    if ai_response_has_violation:
                        ai_policy_names = [v["name"] for v in ai_check_violations]
                        additional_context += f"【AI回答のDDSチェック結果】\n"
                        additional_context += f"- 違反検出: あり\n"
                        additional_context += f"- 違反ポリシー: {', '.join(ai_policy_names)}\n"
                        additional_context += f"- 違反件数: {len(ai_check_violations)}件\n"
                        additional_context += f"- チェックID: {ai_check_request_id}\n"
                    else:
                        additional_context += f"【AI回答のDDSチェック結果】\n"
                        additional_context += f"- 違反検出: なし\n"
                        additional_context += f"- チェックID: {ai_check_request_id}\n"
                else:
                    additional_context += f"【AI回答のDDSチェック】\n"
                    additional_context += f"- 送信がブロックされたため、AI回答はありません\n"
                
                info_result, info_elapsed = run_information_check_ai(
                    file_data,
                    filename,
                    user_input,
                    info_items_enabled,
                    additional_context=additional_context
                )
                
                if info_result:
                    st.success(f"✅ [Inline-2] 情報チェックAI分析完了 (⏱️ {info_elapsed:.2f}秒)")
                    add_debug_log(
                        "Inline-ステップ2",
                        "情報チェックAI分析完了",
                        "success",
                        info_elapsed,
                        {
                            "check_items": [x["name"] for x in info_items_enabled],
                            "response": info_result
                        }
                    )
                    with st.expander("🔎 情報チェックAI分析結果", expanded=True):
                        st.code(info_result, language="text")
                else:
                    st.warning("⚠️ 情報チェックAIの分析結果を取得できませんでした")
                    add_debug_log("Inline-ステップ2", "分析結果を取得できませんでした", "error", info_elapsed)
            
            # ============================================================
            # ステップ3: ZIP作成してDDSに送信（新規）
            # ============================================================
            if info_result and info_result != "検出結果なし":
                step3_start = time.time()
                add_debug_log("Inline-ステップ3", "ZIPファイルを作成してDDSに送信中...", "info")
                st.info("📦 [Inline-3] ZIPファイルを作成してDDSに送信中...")
                
                zip_data = create_zip_for_dds(
                    info_check_result=info_result,
                    file_data=file_data,
                    filename=filename,
                    user_message=user_input,
                    send_target=send_target
                )
                
                zip_size = len(zip_data.getvalue())
                st.info(f"📦 ZIP作成完了: {zip_size/1024:.1f} KB")
                
                v3, request_id3, response_data3, error_info3, elapsed3 = send_zip_to_dds(
                    zip_data=zip_data,
                    dds_url=st.session_state.dds_url,
                    verify_ssl=st.session_state.verify_ssl,
                    source_type="zip",
                    content_block_id="info-check-zip-001"
                )
                
                if v3 is None:
                    v3 = []
                
                st.info(f"✅ [Inline-3] DDS ZIP検査完了 (⏱️ {elapsed3:.2f}秒)")
                add_debug_log("Inline-ステップ3", f"ZIP DDS検査完了", "success", elapsed3, response_data3)
                
                if error_info3:
                    st.error(f"❌ DDSエラー: {error_info3}")
                    add_debug_log("Inline-ステップ3", f"エラー: {error_info3}", "error", elapsed3)
                else:
                    info_dds_response = format_dds_response(response_data3, v3, request_id3) if st.session_state.show_dds_response else None
                    
                    if v3 and len(v3) > 0:
                        policy_names = [v["name"] for v in v3]
                        st.error(f"""
                        🚨 **【警告】情報チェックAI分析結果（ZIP）にポリシー違反が検出されました**
                        
                        情報チェックAIの分析結果と元の内容を含むZIPファイルに以下のポリシー違反が含まれています。
                        
                        **違反ポリシー:** {', '.join(policy_names)}
                        
                        この内容はDLPの補助情報として記録されます。
                        """)
                        add_debug_log("Inline-ステップ4", f"ZIPにポリシー違反: {', '.join(policy_names)}", "error", elapsed3)
                    else:
                        st.success("✅ 情報チェックAI分析結果（ZIP）にポリシー違反はありません")
                        add_debug_log("Inline-ステップ4", "ZIPにポリシー違反なし", "success", elapsed3)
                    
                    zip_contents = []
                    zip_contents.append("📦 **ZIPファイル内容:**")
                    zip_contents.append(f"  - info_check_result.txt: 情報チェックAI分析結果")
                    if user_input and send_target in ["message", "both"]:
                        zip_contents.append(f"  - user_message.txt: ユーザーメッセージ")
                    if file_data and filename and send_target in ["file", "both"]:
                        zip_contents.append(f"  - original_{filename}: 元のファイル")
                    zip_contents.append("")
                    
                    info_msg = f"📊 **情報チェックAI分析結果（ZIPでDDS送信）**\n\n"
                    info_msg += "\n".join(zip_contents)
                    info_msg += "\n"
                    if v3 and len(v3) > 0:
                        policy_names = [v["name"] for v in v3]
                        info_msg += f"🚨 検出された違反: {', '.join(policy_names)}\n\n"
                    else:
                        info_msg += "✅ 違反は検出されませんでした\n\n"
                    info_msg += "---\n"
                    info_msg += f"**分析結果:**\n```\n{info_result}\n```"
                    
                    if st.session_state.blocked_content:
                        info_msg = "🚫 **（送信はブロックされました）**\n\n" + info_msg
                    
                    msg_data = {
                        "role": "assistant",
                        "content": info_msg,
                        "violations": v3,
                        "is_info_check": True
                    }
                    if info_dds_response:
                        msg_data["dds_response"] = info_dds_response
                    st.session_state.messages.append(msg_data)
                    
                    with st.chat_message("assistant"):
                        if st.session_state.blocked_content:
                            st.warning("🚫 送信はブロックされましたが、情報チェックAI分析結果は以下です")
                        st.markdown(info_msg)
                        if info_dds_response:
                            with st.expander("📋 DDSレスポンス詳細", expanded=False):
                                st.markdown(info_dds_response)
            elif info_result == "検出結果なし":
                st.info("ℹ️ 情報チェックAIで検出結果がなかったため、ZIP送信をスキップしました")
            else:
                st.warning("⚠️ 情報チェックAIの結果がないため、ZIP送信をスキップしました")
            
            st.success(f"✅ **Inlineモード完了** (合計: {time.time() - st.session_state.process_start_time:.2f}秒)")
            add_debug_log("Inline-完了", f"全ての処理が完了しました", "success", time.time() - st.session_state.process_start_time)
        
        st.rerun()

# ==================== デバッグパネル ====================
if show_debug and debug_col is not None:
    with debug_col:
        st.subheader("📋 デバッグログ")
        st.caption(f"🕐 最終更新: {datetime.now().strftime('%H:%M:%S')}")
        st.divider()
        
        if st.button("🗑️ ログクリア", key="clear_logs_top_btn", use_container_width=True):
            clear_debug_logs()
            st.rerun()
        
        st.divider()
        
        with st.container(height=550):
            render_debug_logs()
